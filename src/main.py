import os
from pathlib import Path

import requests
import yt_dlp
from openai import OpenAI
from slack_bolt import App
from slack_bolt.adapter.socket_mode import SocketModeHandler
from dotenv import load_dotenv
from docx import Document

# Pour le PDF propre avec retours à la ligne automatiques
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


# 1. CONFIGURATION ET CHARGEMENT
ENV_PATH = Path(__file__).resolve().parents[1] / ".env"
load_dotenv(ENV_PATH)

app = App(token=os.getenv("SLACK_BOT_TOKEN"))
client_openai = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


# --- GÉNÉRATION DE FICHIERS ---
def creer_word(texte, nom_fichier):
    doc = Document()
    doc.add_heading("Transcription Vidéo", 0)
    doc.add_paragraph(texte)
    path = f"{nom_fichier}.docx"
    doc.save(path)
    return path


def creer_pdf(texte, nom_fichier):
    path = f"{nom_fichier}.pdf"
    # Utilisation de Platypus pour gérer les retours à la ligne
    doc = SimpleDocTemplate(path, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    # Titre
    elements.append(Paragraph("Transcription Vidéo", styles["Title"]))
    elements.append(Spacer(1, 12))

    # Corps du texte (on remplace les sauts de ligne pour le format PDF)
    texte_formate = texte.replace("\n", "<br/>")
    elements.append(Paragraph(texte_formate, styles["Normal"]))

    doc.build(elements)
    return path


# --- IA : RÉSUMÉ & TRANSCRIPTION ---
def generer_resume(texte_brut):
    print("🧠 Génération du résumé avec GPT-4o...")
    try:
        reponse = client_openai.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": "Tu es un assistant expert en synthèse. Résume cette transcription de manière professionnelle avec un titre 'Résumé' et une liste 'Les 3 points clés'.",
                },
                {"role": "user", "content": texte_brut},
            ],
        )
        return reponse.choices[0].message.content
    except Exception as e:
        return f"⚠️ Erreur lors du résumé : {str(e)}"


def traiter_audio_et_transcrire(chemin_audio):
    print("🤖 Envoi à OpenAI Whisper...")
    with open(chemin_audio, "rb") as f:
        resultat = client_openai.audio.transcriptions.create(model="whisper-1", file=f)
    return resultat.text


# --- ÉCOUTEURS SLACK ---
@app.event("message")
def handle_message(event, say, client):
    files = event.get("files", [])
    text = event.get("text", "")
    channel = event["channel"]

    # CAS 1 : Vidéo importée du PC (Upload direct dans Slack)
    if files:
        file_info = files[0]
        if file_info["mimetype"].startswith("video/"):
            say("📥 Vidéo PC reçue ! Extraction audio en cours...")
            headers = {"Authorization": f"Bearer {os.getenv('SLACK_BOT_TOKEN')}"}
            r = requests.get(file_info["url_private"], headers=headers)
            with open("temp_video.mp4", "wb") as f:
                f.write(r.content)

            # Conversion via FFmpeg (indispensable sur Ubuntu)
            os.system("ffmpeg -i temp_video.mp4 -q:a 0 -map a audio_temp.mp3 -y")
            transcription = traiter_audio_et_transcrire("audio_temp.mp3")
            envoyer_resultats_complets(client, channel, transcription)

    # CAS 2 : Lien externe (YouTube, Facebook, TikTok)
    elif "http" in text:
        url = text.replace("<", "").replace(">", "").split("|")[0].split("?")[0]
        say("⏳ Lien détecté, téléchargement et analyse...")
        opts = {
            "format": "bestaudio/best",
            "outtmpl": "audio_temp",
            "postprocessors": [
                {
                    "key": "FFmpegExtractAudio",
                    "preferredcodec": "mp3",
                    "preferredquality": "128",
                }
            ],
            "overwrites": True,
        }
        with yt_dlp.YoutubeDL(opts) as ydl:
            ydl.download([url])

        transcription = traiter_audio_et_transcrire("audio_temp.mp3")
        envoyer_resultats_complets(client, channel, transcription)


def envoyer_resultats_complets(client, channel, transcription):
    # 1. Génération du résumé IA
    resume = generer_resume(transcription)
    client.chat_postMessage(channel=channel, text=f"📝 *RÉSUMÉ IA* :\n{resume}")

    # 2. Envoi du texte brut avec boutons d'export
    client.chat_postMessage(
        channel=channel,
        blocks=[
            {
                "type": "section",
                "text": {
                    "type": "mrkdwn",
                    "text": f"✅ *Transcription terminée* :\n\n{transcription[:2800]}...",
                },
            },
            {"type": "divider"},
            {
                "type": "section",
                "text": {
                    "type": "mrkdwn",
                    "text": "*Comment voulez-vous exporter ce document ?*",
                },
            },
            {
                "type": "actions",
                "elements": [
                    {
                        "type": "button",
                        "text": {"type": "plain_text", "text": "Format Word 🟦"},
                        "action_id": "gen_word",
                        "value": transcription[:3000],
                    },
                    {
                        "type": "button",
                        "text": {"type": "plain_text", "text": "Format PDF 🟥"},
                        "action_id": "gen_pdf",
                        "value": transcription[:3000],
                    },
                ],
            },
        ],
    )


# --- ACTIONS DES BOUTONS ---
@app.action("gen_word")
def handle_word(ack, body, client):
    ack()
    texte = body["actions"][0]["value"]
    path = creer_word(texte, "Transcription_Boss")
    client.files_upload_v2(
        channel=body["channel"]["id"], file=path, title="Transcription.docx"
    )


@app.action("gen_pdf")
def handle_pdf(ack, body, client):
    ack()
    texte = body["actions"][0]["value"]
    path = creer_pdf(texte, "Transcription_Boss")
    client.files_upload_v2(
        channel=body["channel"]["id"], file=path, title="Transcription.pdf"
    )


# --- LANCEMENT ---
if __name__ == "__main__":
    app_token = os.getenv("SLACK_APP_TOKEN")
    if not app_token:
        print("❌ ERREUR : SLACK_APP_TOKEN manquant !")
    else:
        print("⚡ Bot en ligne ! Prêt pour les liens et les fichiers PC.")
        SocketModeHandler(app, app_token).start()
